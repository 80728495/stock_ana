"""
Gemini AI 股票分析模块

通过 gemini-webapi（逆向工程）调用 Gemini Web 界面，
利用浏览器 Cookie 认证，无 API 配额限制。
对筛选出的股票进行基本面 + 估值 + 建议的综合分析，
并将结果输出为 .docx 文件。
"""

import asyncio
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from gemini_webapi import GeminiClient
from loguru import logger

from stock_ana.config import OUTPUT_DIR

# ──────── 默认模型 ────────
DEFAULT_MODEL = "gemini-3.0-pro"

# ──────── Prompt 模板 ────────

_PROMPT_TEMPLATE = """角色设定：你是寻求成长股风格的顶级机构分析师。你的目标是寻找具有高成长能力的，但是价格仍然合理不那么离谱的标的。

核心原则：

价格合理：始终评估当前价格与内在价值的偏离度，如果企业显著被高估，并且甚至已经超过了未来一年的完美定价时，那这不该是我们选择的标的。

企业增长：企业出于告诉的行业或者企业增长期，要么仍处于行业发展的中期或者周期的早期，要么虽然行业已经到了中晚期，但是企业所代表的技术却是行业整体要更新换代的方向，企业所拥有的技术在行业中处于中早期。

安全边际：企业过去没有重大的违规或者法律风险，有良好的经营信誉，管理层相对靠谱


任务执行流： 对于股票 {ticker_name}（{ticker_symbol}），执行以下步骤（需调用 Google Search）：

1）基本面分析 (Fundamental)：

检索最新的 10-K/10-Q 财报关键数据（ROE, 自由现金流, 负债率）。

分析商业模式护城河（品牌、转换成本、网络效应）。

估值评估（PE/TTM, PEG, DCF 简算， PS等）。

2）消息面分析 (News/Sentiment)：

搜索过去一段时间在股票分析平台上该企业的新闻，如果发生了重大消息并影响了股价，进一步判断消息对长期基本面的影响是"暂时性"还是"永久性"。


输出方式：
part:1 企业基本面分析和描述，包括商业模式，护城河，当前核心竞争力等。并在段落结尾给企业基本面按照1-10分打分，最高10分。
part:2 企业估值情况，针对其合适的企业当前周期状态给定估值模型。并在段落结尾给估值打1-10分，分数越高代表低估越高。
part:3 结合基本面和估值情况，并给出建议：尽快上车，逐渐买入，暂时不动，逐渐卖出，快速逃跑。"""

# 纳指100常见公司全名映射（部分，其余用 ticker 代替）
_TICKER_NAMES: dict[str, str] = {
    "AAPL": "Apple Inc.",
    "ABNB": "Airbnb Inc.",
    "ADBE": "Adobe Inc.",
    "ADI": "Analog Devices Inc.",
    "ADP": "Automatic Data Processing Inc.",
    "ADSK": "Autodesk Inc.",
    "AEP": "American Electric Power Co.",
    "AMAT": "Applied Materials Inc.",
    "AMD": "Advanced Micro Devices Inc.",
    "AMGN": "Amgen Inc.",
    "AMZN": "Amazon.com Inc.",
    "ANSS": "ANSYS Inc.",
    "APP": "AppLovin Corporation",
    "ARM": "Arm Holdings plc",
    "ASML": "ASML Holding N.V.",
    "AVGO": "Broadcom Inc.",
    "AZN": "AstraZeneca PLC",
    "BIIB": "Biogen Inc.",
    "BKNG": "Booking Holdings Inc.",
    "BKR": "Baker Hughes Company",
    "CCEP": "Coca-Cola Europacific Partners",
    "CDNS": "Cadence Design Systems Inc.",
    "CDW": "CDW Corporation",
    "CEG": "Constellation Energy Group",
    "CHTR": "Charter Communications Inc.",
    "CMCSA": "Comcast Corporation",
    "COST": "Costco Wholesale Corporation",
    "CPRT": "Copart Inc.",
    "CRWD": "CrowdStrike Holdings Inc.",
    "CSCO": "Cisco Systems Inc.",
    "CSGP": "CoStar Group Inc.",
    "CSX": "CSX Corporation",
    "CTAS": "Cintas Corporation",
    "CTSH": "Cognizant Technology Solutions",
    "DASH": "DoorDash Inc.",
    "DDOG": "Datadog Inc.",
    "DLTR": "Dollar Tree Inc.",
    "EA": "Electronic Arts Inc.",
    "EXC": "Exelon Corporation",
    "FANG": "Diamondback Energy Inc.",
    "FAST": "Fastenal Company",
    "FTNT": "Fortinet Inc.",
    "GEHC": "GE HealthCare Technologies",
    "GFS": "GlobalFoundries Inc.",
    "GILD": "Gilead Sciences Inc.",
    "GOOG": "Alphabet Inc.",
    "GOOGL": "Alphabet Inc.",
    "HON": "Honeywell International Inc.",
    "IDXX": "IDEXX Laboratories Inc.",
    "ILMN": "Illumina Inc.",
    "INTC": "Intel Corporation",
    "INTU": "Intuit Inc.",
    "ISRG": "Intuitive Surgical Inc.",
    "KDP": "Keurig Dr Pepper Inc.",
    "KHC": "The Kraft Heinz Company",
    "KLAC": "KLA Corporation",
    "LIN": "Linde plc",
    "LRCX": "Lam Research Corporation",
    "LULU": "Lululemon Athletica Inc.",
    "MAR": "Marriott International Inc.",
    "MCHP": "Microchip Technology Inc.",
    "MDB": "MongoDB Inc.",
    "MDLZ": "Mondelez International Inc.",
    "MELI": "MercadoLibre Inc.",
    "META": "Meta Platforms Inc.",
    "MNST": "Monster Beverage Corporation",
    "MRNA": "Moderna Inc.",
    "MRVL": "Marvell Technology Inc.",
    "MSFT": "Microsoft Corporation",
    "MU": "Micron Technology Inc.",
    "NFLX": "Netflix Inc.",
    "NVDA": "NVIDIA Corporation",
    "NXPI": "NXP Semiconductors N.V.",
    "ODFL": "Old Dominion Freight Line Inc.",
    "ON": "ON Semiconductor Corporation",
    "ORLY": "O'Reilly Automotive Inc.",
    "PANW": "Palo Alto Networks Inc.",
    "PAYX": "Paychex Inc.",
    "PCAR": "PACCAR Inc.",
    "PDD": "PDD Holdings Inc.",
    "PEP": "PepsiCo Inc.",
    "PYPL": "PayPal Holdings Inc.",
    "QCOM": "Qualcomm Inc.",
    "REGN": "Regeneron Pharmaceuticals Inc.",
    "ROP": "Roper Technologies Inc.",
    "ROST": "Ross Stores Inc.",
    "SBUX": "Starbucks Corporation",
    "SMCI": "Super Micro Computer Inc.",
    "SNPS": "Synopsys Inc.",
    "TEAM": "Atlassian Corporation",
    "TMUS": "T-Mobile US Inc.",
    "TSLA": "Tesla Inc.",
    "TTD": "The Trade Desk Inc.",
    "TTWO": "Take-Two Interactive Software",
    "TXN": "Texas Instruments Inc.",
    "VRSK": "Verisk Analytics Inc.",
    "VRTX": "Vertex Pharmaceuticals Inc.",
    "WBD": "Warner Bros. Discovery Inc.",
    "WDAY": "Workday Inc.",
    "XEL": "Xcel Energy Inc.",
    "ZS": "Zscaler Inc.",
}


def _get_ticker_name(ticker: str) -> str:
    """获取股票全名，未知的返回 ticker 本身"""
    return _TICKER_NAMES.get(ticker.upper(), ticker)


def _build_prompt(ticker: str) -> str:
    """为指定股票构造分析 prompt"""
    name = _get_ticker_name(ticker)
    return _PROMPT_TEMPLATE.format(ticker_name=name, ticker_symbol=ticker)


# ──────── 客户端初始化 ────────

async def _init_client(model: str = DEFAULT_MODEL) -> GeminiClient:
    """
    初始化 Gemini Web 客户端。

    依赖 browser-cookie3 自动从 Chrome 浏览器读取 Cookie，
    需要先在 Chrome 中登录 https://gemini.google.com 。
    """
    client = GeminiClient()
    await client.init(
        timeout=120,
        auto_close=False,
        auto_refresh=True,
        verbose=False,  # 减少初始化日志（如 Safari 权限警告）
    )
    logger.info(f"Gemini Web 客户端初始化成功（模型：{model}）")
    return client


# ──────── 核心分析 ────────

async def analyze_stock(
    ticker: str,
    client: GeminiClient | None = None,
    model: str = DEFAULT_MODEL,
    max_retries: int = 2,
) -> str:
    """
    调用 Gemini Web 分析单只股票。

    Args:
        ticker: 股票代号（如 NVDA）
        client: GeminiClient 实例，为 None 则自动创建
        model: 使用的模型名
        max_retries: 遇到错误时的最大重试次数

    Returns:
        Gemini 返回的分析文本
    """
    own_client = client is None
    if own_client:
        client = await _init_client(model)

    prompt = _build_prompt(ticker)
    name = _get_ticker_name(ticker)
    logger.info(f"正在分析 {ticker} ({name})...")

    last_err = None
    for attempt in range(max_retries + 1):
        try:
            response = await client.generate_content(prompt, model=model)
            text = response.text or ""
            logger.success(f"✅ {ticker} 分析完成，共 {len(text)} 字符")
            return text
        except Exception as e:
            last_err = e
            err_str = str(e)
            if attempt < max_retries:
                wait = 10 * (attempt + 1)
                logger.warning(
                    f"⏳ {ticker} 请求失败: {err_str[:100]}，"
                    f"等待 {wait}s 后重试 ({attempt + 1}/{max_retries})..."
                )
                await asyncio.sleep(wait)
            else:
                logger.error(
                    f"❌ {ticker} 分析失败（已重试 {max_retries} 次）: {err_str[:200]}"
                )

    if own_client:
        await client.close()

    raise last_err  # type: ignore[misc]


# ──────── docx 生成 ────────

def _save_to_docx(ticker: str, analysis_text: str, output_dir: Path) -> Path:
    """
    将分析结果保存为 .docx 文件。
    """
    doc = Document()

    # ── 标题 ──
    name = _get_ticker_name(ticker)
    title = doc.add_heading(f"{ticker} - {name} 投资分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 副标题（日期） ──
    from datetime import date

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"分析日期：{date.today().isoformat()}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("")  # 空行

    # ── 正文：按 part 分段 ──
    parts = re.split(r"(?i)\n*\s*part\s*[:：]\s*", analysis_text)

    for part in parts:
        part = part.strip()
        if not part:
            continue

        lines = part.split("\n")
        first_line = lines[0].strip()

        heading_match = re.match(r"^(\d+)\s*[.、)）]?\s*(.*)", first_line)
        if heading_match:
            part_num = heading_match.group(1)
            part_title = heading_match.group(2) or first_line
            doc.add_heading(f"Part {part_num}: {part_title}", level=1)
            body_lines = lines[1:]
        else:
            body_lines = lines

        body_text = "\n".join(body_lines).strip()
        if body_text:
            paragraphs = body_text.split("\n\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                if re.match(r"^[\-\*•]", para_text):
                    for item in para_text.split("\n"):
                        item = item.strip().lstrip("-*• ")
                        if item:
                            doc.add_paragraph(item, style="List Bullet")
                elif re.match(r"^\d+[\.\)）、]", para_text):
                    for item in para_text.split("\n"):
                        item = item.strip()
                        if item:
                            doc.add_paragraph(item, style="List Number")
                else:
                    for sub_para in para_text.split("\n"):
                        sub_para = sub_para.strip()
                        if sub_para:
                            p = doc.add_paragraph()
                            _add_rich_text(p, sub_para)

    # ── 保存 ──
    output_dir.mkdir(parents=True, exist_ok=True)
    save_path = output_dir / f"{ticker}_analysis.docx"
    doc.save(str(save_path))
    return save_path


def _add_rich_text(paragraph, text: str):
    """解析 Markdown 粗体 **text** 并添加到段落"""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# ──────── 组合接口 ────────

async def analyze_and_save(
    ticker: str,
    client: GeminiClient | None = None,
    model: str = DEFAULT_MODEL,
    output_dir: Path | None = None,
) -> Path:
    """分析单只股票并保存为 docx。"""
    if output_dir is None:
        output_dir = OUTPUT_DIR

    text = await analyze_stock(ticker, client=client, model=model)
    path = _save_to_docx(ticker, text, output_dir)
    logger.info(f"📄 {ticker} 分析报告已保存 → {path}")
    return path


async def batch_analyze(
    tickers: list[str],
    output_dir: Path | None = None,
    model: str = DEFAULT_MODEL,
    delay: float = 5.0,
) -> list[Path]:
    """
    批量分析多只股票。

    Args:
        tickers: 股票代号列表
        output_dir: 输出目录
        model: Gemini 模型名
        delay: 每次调用间的延迟（秒），避免限流

    Returns:
        生成的 docx 文件路径列表
    """
    if output_dir is None:
        output_dir = OUTPUT_DIR

    client = await _init_client(model)
    results: list[Path] = []

    logger.info(f"开始批量分析 {len(tickers)} 只股票（模型：{model}）...")

    try:
        for i, ticker in enumerate(tickers, 1):
            try:
                logger.info(f"[{i}/{len(tickers)}] 分析 {ticker}...")
                path = await analyze_and_save(
                    ticker, client=client, model=model, output_dir=output_dir
                )
                results.append(path)

                if i < len(tickers) and delay > 0:
                    logger.info(f"⏳ 等待 {delay:.0f}s...")
                    await asyncio.sleep(delay)

            except Exception as e:
                logger.error(f"❌ {ticker} 分析失败: {e}")
                continue
    finally:
        await client.close()

    logger.info(f"批量分析完成：{len(results)}/{len(tickers)} 只成功")
    return results


async def analyze_screener_results(
    hits: list[dict],
    output_dir: Path | None = None,
    model: str = DEFAULT_MODEL,
    delay: float = 5.0,
) -> list[Path]:
    """
    对筛选结果列表进行批量分析。

    Args:
        hits: 筛选结果，每个元素有 "ticker" 字段
        output_dir: 输出目录
        model: Gemini 模型名
        delay: 调用间延迟

    Returns:
        生成的 docx 文件路径列表
    """
    tickers = list(dict.fromkeys(h["ticker"] for h in hits))  # 去重保序
    return await batch_analyze(tickers, output_dir=output_dir, model=model, delay=delay)


# ──────── 第三步：综合排序 ────────

_RANK_PROMPT = """你是一位顶级机构投资组合经理。我上传了多份个股分析报告（docx 文件），每份报告包含：
- Part 1：基本面分析（含 1-10 分评分）
- Part 2：估值分析（含 1-10 分评分，分数越高代表越低估）
- Part 3：操作建议

请你完成以下任务：

1）仔细阅读每份报告，提取每只股票的：
   - 基本面评分
   - 估值评分
   - 操作建议

2）计算综合评分：综合评分 = 基本面评分 × 0.4 + 估值评分 × 0.6
   （我们更看重估值的安全边际）

3）按综合评分从高到低排序，输出排名表格，格式如下：

排名 | 股票代号 | 公司名称 | 基本面评分 | 估值评分 | 综合评分 | 操作建议
1    | XXXX    | xxx     | x/10      | x/10    | x.x    | 逐渐买入

4）最后给出你的投资组合建议：
   - 首选标的（最值得优先配置的 1-3 只）
   - 备选标的（可以关注但不急于建仓的）
   - 规避标的（建议暂时回避的）

请确保排名表格完整、评分准确。"""


async def rank_and_summarize(
    report_paths: list[Path] | None = None,
    output_dir: Path | None = None,
    model: str = DEFAULT_MODEL,
) -> Path:
    """
    第三步：将所有个股分析报告上传给 Gemini，综合排序。

    Args:
        report_paths: docx 报告路径列表。为 None 则自动扫描 output_dir 中的 *_analysis.docx
        output_dir: 输出目录
        model: Gemini 模型名

    Returns:
        排序汇总报告路径
    """
    if output_dir is None:
        output_dir = OUTPUT_DIR

    # 收集报告文件
    if report_paths is None:
        report_paths = sorted(output_dir.glob("*_analysis.docx"))

    if not report_paths:
        raise ValueError(f"未找到分析报告！请先运行第二步生成个股报告到 {output_dir}")

    file_names = [p.name for p in report_paths]
    logger.info(f"📊 开始综合排序，共 {len(report_paths)} 份报告：{file_names}")

    # 初始化客户端
    client = await _init_client(model)

    try:
        # 上传文件 + prompt
        file_strs = [str(p) for p in report_paths]
        response = await client.generate_content(
            _RANK_PROMPT,
            files=file_strs,
            model=model,
        )
        text = response.text or ""
        logger.success(f"✅ 综合排序完成，共 {len(text)} 字符")

        # 保存为汇总 docx
        path = _save_rank_docx(text, report_paths, output_dir)
        logger.info(f"📄 综合排序报告已保存 → {path}")
        return path

    finally:
        await client.close()


def _save_rank_docx(
    rank_text: str, report_paths: list[Path], output_dir: Path
) -> Path:
    """将排序结果保存为汇总 docx。"""
    doc = Document()

    # ── 标题 ──
    title = doc.add_heading("投资标的综合排序报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 副标题 ──
    from datetime import date

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"分析日期：{date.today().isoformat()}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 分析标的列表
    tickers = [p.stem.replace("_analysis", "") for p in report_paths]
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub2.add_run(f"涵盖标的：{', '.join(tickers)}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")

    # ── 正文 ──
    for line in rank_text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Markdown 标题
        heading_match = re.match(r"^(#{1,3})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2), level=level)
            continue

        # 表格行（| col1 | col2 | ...）
        if line.startswith("|") and "|" in line[1:]:
            # 跳过分隔行 |---|---|
            if re.match(r"^\|[\s\-:]+\|", line):
                continue
            # 这里简单用段落呈现表格行
            p = doc.add_paragraph()
            _add_rich_text(p, line)
            continue

        # 列表项
        if re.match(r"^[\-\*•]", line):
            item = line.lstrip("-*• ").strip()
            if item:
                doc.add_paragraph(item, style="List Bullet")
            continue

        if re.match(r"^\d+[\.\)）、]", line):
            doc.add_paragraph(line, style="List Number")
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_rich_text(p, line)

    # ── 保存 ──
    output_dir.mkdir(parents=True, exist_ok=True)
    save_path = output_dir / "00_综合排序报告.docx"
    doc.save(str(save_path))
    return save_path
